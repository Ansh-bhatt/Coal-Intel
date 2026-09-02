"""Text extraction for uploaded documents.

Extracts per-page text (PDF) or structured key-value rows (XLSX/DOCX) and
turns them into raw ``ExtractedRecord`` rows with confidence scores.

PDF text extraction uses the embedded text layer (pypdf).  Scanned PDFs
would need an OCR step — the ``extract_pdf_text`` function is the single
swap point for a real OCR/document-AI service, keeping the abstraction
required by the PRD.
"""

import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook

from app.core.config import get_settings

settings = get_settings()

STOPWORDS = {
    "the", "and", "for", "with", "from", "this", "that", "are", "was",
    "were", "been", "will", "have", "has", "had", "not", "but", "its",
}

@dataclass
class ExtractedPage:
    """Text layer for a single page, plus a confidence estimate."""

    page_number: int
    text: str
    confidence: float


@dataclass
class ExtractionResult:
    pages: list[ExtractedPage] = field(default_factory=list)
    records: list[dict] = field(default_factory=list)


def _estimate_confidence(text: str) -> float:
    """Heuristic confidence for a text layer (0.0 – 1.0)."""
    if not text or not text.strip():
        return 0.1
    words = re.findall(r"[A-Za-z0-9]+", text)
    if not words:
        return 0.1
    ascii_words = sum(1 for w in words if w.isascii())
    ratio = ascii_words / len(words)
    weird = len(re.findall(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", text))
    penalty = min(0.3, weird * 0.02)
    return round(max(0.0, min(0.99, ratio * 0.99 - penalty)), 3)


def extract_pdf_text(path: Path) -> list[ExtractedPage]:
    """Extract per-page text from a PDF's embedded text layer."""
    reader = PdfReader(str(path))
    pages: list[ExtractedPage] = []
    for idx, page in enumerate(reader.pages, start=1):
        raw = page.extract_text() or ""
        pages.append(
            ExtractedPage(
                page_number=idx,
                text=raw.strip(),
                confidence=_estimate_confidence(raw),
            )
        )
    return pages


def _extract_records_from_page(page: ExtractedPage) -> list[dict]:
    """Heuristic key-value extraction from a text page.

    Looks for ``key : value`` lines — the shape of the seeded demo reports.
    Values that fail to parse cleanly get a low confidence so they surface in
    the HITL grid.
    """
    records: list[dict] = []
    for line in page.text.splitlines():
        line = line.strip()
        if not line or len(line) < 5:
            continue
        if line.isupper() and len(line.split()) <= 4:
            continue
        match = re.match(r"^([A-Za-z][A-Za-z0-9 ()&/.-]{3,80}?)[:\-]\s+(.+)$", line)
        if not match:
            continue
        key, value = match.group(1).strip(), match.group(2).strip()
        if not key or not value or value.lower() in STOPWORDS:
            continue
        confidence = page.confidence
        if not re.search(r"[0-9%]|MT|Mcum|INR|Cr|₹|million|tonnes?$", value, re.I):
            confidence = max(0.1, round(confidence - 0.35, 3))
        records.append(
            {
                "id": str(uuid.uuid4()),
                "key": key,
                "value": value,
                "confidence": confidence,
            }
        )
    return records

def _extract_records_from_workbook(path: Path) -> list[dict]:
    """Extract key-value pairs from the first worksheet of an XLSX."""
    wb = load_workbook(str(path), data_only=True, read_only=True)
    ws = wb.active
    records: list[dict] = []
    for row in ws.iter_rows(values_only=True):
        if not row:
            continue
        cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
        if len(cells) < 2:
            continue
        key, value = cells[0], cells[1]
        records.append(
            {
                "id": str(uuid.uuid4()),
                "key": key,
                "value": value,
                "confidence": 0.97,
            }
        )
    wb.close()
    return records


def _extract_records_from_docx(path: Path) -> list[dict]:
    """Extract key-value pairs from a DOCX using its paragraphs/tables."""
    doc = DocxDocument(str(path))
    records: list[dict] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or ":" not in text:
            continue
        key, value = text.split(":", 1)
        key, value = key.strip(), value.strip()
        if key and value:
            records.append(
                {
                    "id": str(uuid.uuid4()),
                    "key": key,
                    "value": value,
                    "confidence": 0.95,
                }
            )
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if len(cells) >= 2:
                records.append(
                    {
                        "id": str(uuid.uuid4()),
                        "key": cells[0],
                        "value": cells[1],
                        "confidence": 0.97,
                    }
                )
    return records


def extract_document(path: Path, file_type: str) -> ExtractionResult:
    """Dispatch to the right extractor based on file type."""
    file_type = file_type.lower()
    if file_type == "pdf":
        pages = extract_pdf_text(path)
        records: list[dict] = []
        for page in pages:
            records.extend(_extract_records_from_page(page))
        return ExtractionResult(pages=pages, records=records)
    if file_type in ("xlsx", "xls"):
        return ExtractionResult(pages=[], records=_extract_records_from_workbook(path))
    if file_type in ("docx", "doc"):
        return ExtractionResult(pages=[], records=_extract_records_from_docx(path))
    raise ValueError(f"Unsupported file type: {file_type}")


def compute_status(confidence: float) -> str:
    """Apply the HITL rule: < threshold → flagged, else pending (auto-verified)."""
    if confidence < settings.low_confidence_threshold:
        return "flagged"
    return "pending"